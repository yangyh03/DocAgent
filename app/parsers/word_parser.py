from __future__ import annotations

import mimetypes
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

# Word 图片节点里的 r:embed / r:link 使用这个关系命名空间。
REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
TEXT_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MATH_TAG = f"{{{MATH_NS}}}"
FORMULA_STYLE_NODES = {"ctrlPr", "rPr", "fPr", "sSupPr", "sSubPr", "naryPr"}


def parse_word_document(
    file_path: str | Path,
    assets_dir: str | Path | None = None,
) -> dict[str, Any]:
    """按 DOCX 原始顺序解析 Word 文档。

    输出两层结果：
    1. blocks：面向结构化处理的有序块，类型为 paragraph / title / table / image。
    2. content：由 blocks 拼接出来的兼容文本，供接口 fileContent 使用。
    """
    path = Path(file_path)
    asset_root = Path(assets_dir) if assets_dir else path.parent / "assets"
    asset_root.mkdir(parents=True, exist_ok=True)

    doc = Document(path)
    blocks: list[dict[str, Any]] = []
    assets: list[dict[str, str]] = []

    # body.iterchildren() 是从 Word XML 的 body 里按原始顺序遍历元素，能保留真实顺序
    image_index = 0    # 图片计数命名
    formula_index = 0   # 公式计数命名
    paragraph_index = 0
    table_index = 0
    has_seen_nonempty_paragraph = False   # 判断当前段落是不是文档第一个非空段落，用来做标题兜底识别

    for element in doc.element.body.iterchildren():
        if element.tag.endswith("}p"):
            paragraph = doc.paragraphs[paragraph_index]
            paragraph_index += 1
            paragraph_text = _paragraph_text_for_classification(paragraph)
            is_first_nonempty = bool(paragraph_text) and not has_seen_nonempty_paragraph
            image_index, formula_index = _append_paragraph_blocks(
                paragraph=paragraph,
                blocks=blocks,
                assets=assets,
                asset_root=asset_root,
                image_index=image_index,
                formula_index=formula_index,
                is_first_nonempty=is_first_nonempty,
            )
            if paragraph_text:
                has_seen_nonempty_paragraph = True
        elif element.tag.endswith("}tbl"):
            table = doc.tables[table_index]
            table_index += 1
            blocks.append(_table_to_block(table, table_index))

    return {
        "content": render_blocks_to_text(blocks),
        "status": "success",
        "blocks": blocks,
        "assets": assets,
    }


def render_blocks_to_text(blocks: list[dict[str, Any]]) -> str:
    """把结构化 blocks 渲染成一份可读文本，兼容 fileContent。"""
    lines: list[str] = []
    for block in blocks:
        block_type = block["type"]
        if block_type in {"title", "paragraph"}:
            lines.append(block["content"])
        elif block_type == "table" and block.get("content"):
            lines.append(block["content"])
        elif block_type == "image":
            metadata = block.get("metadata", {})
            file_name = metadata.get("file_name", "unknown")
            description = (block.get("content") or metadata.get("description") or "待生成图片描述").strip()
            lines.append(f"[图片: {file_name} 图片内容: {description}]")
    return "\n".join(line for line in lines if line.strip())


def _append_paragraph_blocks(
    paragraph: Paragraph,
    blocks: list[dict[str, Any]],
    assets: list[dict[str, str]],
    asset_root: Path,
    image_index: int,
    formula_index: int,
    is_first_nonempty: bool,
) -> tuple[int, int]:
    """解析一个段落，并把段落中的图片、公式插入到原始出现位置。"""
    text_parts: list[str] = []
    paragraph_metadata = _paragraph_kind_and_metadata(
        paragraph=paragraph,
        text=_paragraph_text_for_classification(paragraph),
        is_first_nonempty=is_first_nonempty,
    )
    paragraph_kind, metadata = paragraph_metadata
    formulas: list[dict[str, Any]] = []

    def flush_text() -> None:
        """遇到图片前先落盘已有文本，避免图片被挪到段落末尾。"""
        text = "".join(text_parts).strip()
        text_parts.clear()
        if text:
            block_metadata = dict(metadata)
            if formulas:
                block_metadata["formulas"] = list(formulas)
            blocks.append(
                {
                    "type": paragraph_kind,
                    "content": text,
                    "metadata": block_metadata,
                }
            )
            formulas.clear()

    # 直接按段落 XML 子节点扫描，才能同时保留文本、图片和 Office Math 公式的相对顺序。
    for child in paragraph._element.iterchildren():
        if _is_formula_node(child):
            formula_index += 1
            placeholder = f"[公式{formula_index}]"
            formula_text = _formula_to_readable(child)
            text_parts.append(placeholder)
            formulas.append(
                {
                    "formula_index": formula_index,
                    "position_hint": placeholder,
                    "readable_text": formula_text["readable_text"],
                    "latex_text": formula_text["latex_text"],
                }
            )
            continue

        text_parts.append(_extract_text_from_element(child))
        image_refs = _extract_image_refs_from_element(child)
        if not image_refs:
            continue

        flush_text()
        for r_id in image_refs:
            image_index += 1
            image_block = _save_image_block(
                paragraph=paragraph,
                r_id=r_id,
                asset_root=asset_root,
                assets=assets,
                image_index=image_index,
            )
            if image_block:
                blocks.append(image_block)

    flush_text()
    return image_index, formula_index


def _paragraph_kind_and_metadata(
    paragraph: Paragraph,
    text: str,
    is_first_nonempty: bool,
) -> tuple[str, dict[str, Any]]:
    """根据 Word 样式和保守启发式判断普通段落还是标题。"""
    style_name = paragraph.style.name if paragraph.style is not None else ""
    metadata: dict[str, Any] = {"style": style_name}

    if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        metadata["alignment"] = "center"

    title_level = _title_level(style_name)
    if title_level is not None:
        metadata["level"] = title_level
        metadata["title_source"] = "style"
        return "title", metadata

    numbered_level = _numbered_heading_level(text)
    if numbered_level is not None:
        metadata["level"] = numbered_level
        metadata["title_source"] = "numbered_heading"
        return "title", metadata

    if is_first_nonempty and _looks_like_document_title(text):
        metadata["level"] = 1
        metadata["title_source"] = "first_paragraph"
        return "title", metadata

    return "paragraph", metadata


def _title_level(style_name: str) -> int | None:
    """兼容英文 Heading 1 和中文 标题 1 这两类常见标题样式。"""
    normalized = style_name.strip().lower()
    if normalized.startswith("heading") or normalized.startswith("标题"):
        match = re.search(r"(\d+)", normalized)
        return int(match.group(1)) if match else 1
    return None


def _numbered_heading_level(text: str) -> int | None:
    """识别 1. 引言、1.1 模型结构、第1章 绪论、一、背景 这类常见标题。"""
    normalized = text.strip()
    if not normalized or len(normalized) > 80:
        return None

    decimal_match = re.match(r"^(\d+(?:\.\d+)*)(?:[\.、．]\s*|\s+)\S+", normalized)
    if decimal_match and "." in decimal_match.group(1):
        return decimal_match.group(1).count(".") + 1
    if re.match(r"^\d+[\.、．]\s*\S+", normalized):
        return 1
    if re.match(r"^第[一二三四五六七八九十\d]+[章节部分]\s*\S*", normalized):
        return 1
    if re.match(r"^[一二三四五六七八九十]+[、.．]\s*\S+", normalized):
        return 1
    return None


def _looks_like_document_title(text: str) -> bool:
    """首个非空段落的兜底标题识别，规则保持保守，避免把长正文误判为标题。"""
    normalized = text.strip()
    if not (4 <= len(normalized) <= 80):
        return False
    if normalized.endswith(("。", "！", "？", "；", ".", "!", "?", ";")):
        return False
    return True


def _extract_image_refs_from_element(element) -> list[str]:
    """从 OOXML 节点中提取图片 relationship id。"""
    refs: list[str] = []
    for node in element.iter():
        if not node.tag.endswith("}blip"):
            continue
        r_id = node.get(f"{{{REL_NS}}}embed") or node.get(f"{{{REL_NS}}}link")
        if r_id:
            refs.append(r_id)
    return refs


def _extract_text_from_element(element) -> str:
    """提取普通 Word 文本节点，跳过公式中的 m:t，避免公式文本和占位符重复。"""
    texts: list[str] = []
    for node in element.iter():
        if node.tag == f"{{{TEXT_NS}}}t" and node.text:
            texts.append(node.text)
    return "".join(texts)


def _paragraph_text_for_classification(paragraph: Paragraph) -> str:
    """给标题启发式使用的纯文本，不包含公式 XML。"""
    return "".join(_extract_text_from_element(child) for child in paragraph._element.iterchildren()).strip()


def _is_formula_node(element) -> bool:
    return element.tag.endswith("}oMath") or element.tag.endswith("}oMathPara")


def _formula_to_readable(element) -> dict[str, str]:
    """把常见 OMML 公式结构转换成可读文本和简化 LaTeX。

    这是轻量解析器，不追求覆盖全部 Word 公式语法；遇到未知结构时降级拼接 m:t 文本。
    """
    readable_text, latex_text = _render_formula_element(element)
    fallback = _extract_math_text(element)
    readable_text = _clean_formula_text(readable_text or fallback)
    latex_text = _clean_formula_text(latex_text or _latex_text(fallback))
    return {
        "readable_text": readable_text,
        "latex_text": latex_text,
    }


def _render_formula_element(element) -> tuple[str, str]:
    if element is None or not isinstance(element.tag, str):
        return "", ""

    name = _local_name(element)
    if name in FORMULA_STYLE_NODES:
        return "", ""
    if name in {"oMathPara", "oMath", "e", "num", "den", "sub", "sup"}:
        return _render_formula_children(element)
    if name == "r":
        text = _extract_math_text(element)
        return text, _latex_text(text)
    if name == "f":
        numerator = _render_formula_part(element, "num")
        denominator = _render_formula_part(element, "den")
        return (
            f"({numerator[0]})/({denominator[0]})",
            f"\\frac{{{numerator[1]}}}{{{denominator[1]}}}",
        )
    if name == "sSup":
        base = _render_formula_part(element, "e")
        superscript = _render_formula_part(element, "sup")
        return (
            f"{base[0]}^({superscript[0]})",
            f"{base[1]}^{{{superscript[1]}}}",
        )
    if name == "sSub":
        base = _render_formula_part(element, "e")
        subscript = _render_formula_part(element, "sub")
        return (
            f"{base[0]}_({subscript[0]})",
            f"{base[1]}_{{{subscript[1]}}}",
        )
    if name == "nary":
        return _render_nary_formula(element)

    return _render_formula_children(element)


def _render_nary_formula(element) -> tuple[str, str]:
    operator = _nary_operator(element)
    readable_operator = "sum" if operator == "∑" else operator
    latex_operator = r"\sum" if operator == "∑" else _latex_text(operator)
    subscript = _render_formula_part(element, "sub")
    superscript = _render_formula_part(element, "sup")
    body = _render_formula_part(element, "e")

    readable = f"{readable_operator}_({subscript[0]})^({superscript[0]})"
    latex = f"{latex_operator}_{{{subscript[1]}}}^{{{superscript[1]}}}"
    if body[0]:
        readable = f"{readable} {body[0]}"
    if body[1]:
        latex = f"{latex} {body[1]}"
    return readable, latex


def _render_formula_part(element, child_name: str) -> tuple[str, str]:
    child = element.find(f"{MATH_TAG}{child_name}")
    return _render_formula_element(child)


def _render_formula_children(element) -> tuple[str, str]:
    readable_parts: list[str] = []
    latex_parts: list[str] = []
    for child in element:
        if not isinstance(child.tag, str) or _local_name(child) in FORMULA_STYLE_NODES:
            continue
        readable, latex = _render_formula_element(child)
        readable_parts.append(readable)
        latex_parts.append(latex)
    return "".join(readable_parts), "".join(latex_parts)


def _nary_operator(element) -> str:
    properties = element.find(f"{MATH_TAG}naryPr")
    if properties is None:
        return "∑"
    operator = properties.find(f"{MATH_TAG}chr")
    if operator is None:
        return "∑"
    return operator.get(f"{MATH_TAG}val") or operator.get("val") or "∑"


def _extract_math_text(element) -> str:
    texts: list[str] = []
    for node in element.iter():
        if node.tag == f"{MATH_TAG}t" and node.text:
            texts.append(node.text)
    return "".join(texts)


def _latex_text(text: str) -> str:
    return text.replace("|", r" \mid ")


def _clean_formula_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u200a", "")).strip()


def _local_name(element) -> str:
    return etree.QName(element).localname


def _save_image_block(
    paragraph: Paragraph,
    r_id: str,
    asset_root: Path,
    assets: list[dict[str, str]],
    image_index: int,
) -> dict[str, Any] | None:
    """通过 relationship id 找到图片二进制，保存到 assets 并返回 image block。"""
    image_part = paragraph.part.related_parts.get(r_id)
    if image_part is None:
        return None

    mime_type = getattr(image_part, "content_type", None) or "application/octet-stream"
    extension = mimetypes.guess_extension(mime_type) or Path(str(image_part.partname)).suffix or ".bin"
    file_name = f"image_{image_index}{extension}"
    target_path = asset_root / file_name
    target_path.write_bytes(image_part.blob)

    asset = {
        "file_name": file_name,
        "path": str(target_path),
        "mime_type": mime_type,
    }
    assets.append(asset)

    return {
        "type": "image",
        "content": "",
        "metadata": asset,
    }


def _table_to_block(table: Table, table_index: int) -> dict[str, Any]:
    """只解析一级表格：保留行列文本，不递归嵌套表格。"""
    rows: list[list[str]] = []
    for row in table.rows:
        row_values: list[str] = []
        for cell in row.cells:
            cell_text = "\n".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            row_values.append(cell_text)
        rows.append(row_values)

    return {
        "type": "table",
        "content": _rows_to_markdown(rows),
        "metadata": {
            "table_index": table_index,
            "rows": rows,
            "row_count": len(rows),
            "column_count": max((len(row) for row in rows), default=0),
        },
    }


def _rows_to_markdown(rows: list[list[str]]) -> str:
    """给 fileContent 使用的简洁表格文本，真实结构仍以 metadata.rows 为准。"""
    if not rows:
        return ""
    return "\n".join("| " + " | ".join(_clean_cell(cell) for cell in row) + " |" for row in rows)


def _clean_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\r", " ").replace("\n", " ").strip()
